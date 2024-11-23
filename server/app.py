from flask import Flask, request, send_file, jsonify
from werkzeug.utils import secure_filename
from fpdf import FPDF
from PyPDF2 import PdfWriter
import os
from flask_cors import CORS
import docx
from io import BytesIO

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = 'uploaded_files'
OUTPUT_FOLDER = 'converted_files'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Function to extract text from DOCX
def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    content = ""
    for para in doc.paragraphs:
        content += para.text + "\n"
    return content

# Function to extract images from DOCX and return as list of BytesIO objects
def extract_images_from_docx(docx_path):
    doc = docx.Document(docx_path)
    images = []
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            img_data = rel.target_part.blob
            img_io = BytesIO(img_data)
            images.append(img_io)
    return images

@app.route('/docxtopdf', methods=['POST'])
def convert_docx_to_pdf():
    if 'docFile' not in request.files:
        return jsonify({"message": "No file uploaded"}), 400

    doc_file = request.files['docFile']
    password = request.form.get('password', None)  # Optional password
    if doc_file.filename == '':
        return jsonify({"message": "No selected file"}), 400

    # Save the uploaded file
    filename = secure_filename(doc_file.filename)
    input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    doc_file.save(input_path)

    # Generate the output PDF filename
    output_filename = f"{os.path.splitext(filename)[0]}.pdf"
    output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)

    try:
        # Extract text and images from DOCX
        content = extract_text_from_docx(input_path)
        images = extract_images_from_docx(input_path)
        
        # Convert DOCX to PDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)

        # Write the DOCX content into the PDF
        for line in content.split('\n'):
            pdf.cell(200, 10, txt=line, ln=True, align="L")
        
        # Add images to the PDF
        for img in images:
            pdf.add_page()
            img_path = os.path.join(app.config['OUTPUT_FOLDER'], "temp_image.jpg")
            with open(img_path, 'wb') as f:
                f.write(img.getvalue())
            pdf.image(img_path, x=10, y=10, w=100)  # Adjust size and position accordingly

        # Output the PDF
        pdf.output(output_path)

        # Apply password if provided
        if password:
            writer = PdfWriter()
            with open(output_path, "rb") as original_pdf:
                writer.add_page(writer.pages[0])  # Add all pages
                writer.encrypt(user_pwd=password, owner_pwd=password)
            with open(output_path, "wb") as encrypted_pdf:
                writer.write(encrypted_pdf)

        # Send the file as a response
        return send_file(output_path, as_attachment=True)
    
    except Exception as e:
        return jsonify({"message": f"Error during conversion: {str(e)}"}), 500
    
    finally:
        # Cleanup uploaded file and temporary images
        if os.path.exists(input_path):
            os.remove(input_path)
        for img in images:
            img.close()

if __name__ == '__main__':
    app.run(port=3000, debug=True)
